# -*- coding: utf-8 -*-
"""
generate_report.py
===================
YMH340 Veri Madenciligi Donem Projesi — NIHAI RAPOR URETICISI

Bu betik, ana hat (main.py) tarafindan uretilen tablolari (outputs/tables) ve
gorselleri (outputs/figures) okuyarak, proje kilavuzunda istenen formata
birebir uyan bir Word (.docx) raporu olusturur:

  * Kapak sayfasi (proje basligi + ekip uyeleri)
  * 6 ana bolum (her adim ayri bolum halinde)
  * IEEE numarali referanslar
  * Tum kenar bosluklari 2.5 cm, 12 punto Times New Roman, 1.5 satir araligi,
    iki yana yasli hizalama

Cikti: report/YMH340_Proje_Raporu.docx
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --------------------------------------------------------------------------- #
# YOLLAR
# --------------------------------------------------------------------------- #
ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "outputs" / "figures"
TAB = ROOT / "outputs" / "tables"
OUT = ROOT / "report" / "YMH340_Proje_Raporu.docx"

# Renkler
BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x22, 0x22, 0x22)
HEAD_FILL = "1F4E79"
ALT_FILL = "DCE6F1"

FIG_W = Cm(15.5)   # tam genislik figur (A4 icerik genisligi ~16 cm)


# --------------------------------------------------------------------------- #
# BICIMLENDIRME YARDIMCILARI
# --------------------------------------------------------------------------- #
def set_cell_background(cell, hex_color: str) -> None:
    """Bir tablo hucresine arka plan rengi uygular."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_repeat_header(row) -> None:
    """Tablo baslik satirinin her sayfada tekrarlanmasini saglar."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_page_numbers(section) -> None:
    """Alt bilgiye 'Sayfa X' alanini (otomatik sayfa numarasi) ekler."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sayfa ")
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"
    fld1 = OxmlElement("w:fldSimple")
    fld1.set(qn("w:instr"), "PAGE")
    p._p.append(fld1)


def style_normal(doc: Document) -> None:
    """Belge geneli icin Times New Roman 12 punto, 1.5 satir araligi tanimlar."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    # Dogu Asya fontunu da ayarla (uyumluluk icin)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)


def para(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
         italic=False, size=12, color=None, space_after=6, space_before=0,
         first_line=None, line_spacing=1.5):
    """Bicimlendirilmis bir paragraf ekler."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = line_spacing
    if first_line:
        p.paragraph_format.first_line_indent = Cm(first_line)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color is not None:
        run.font.color.rgb = color
    return p


def heading(doc, text, level=1):
    """Numarali bolum basligi (Times New Roman, koyu)."""
    sizes = {1: 14, 2: 12.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = "Times New Roman"
    run.font.color.rgb = BLUE
    return p


def bullet(doc, text, bold_lead=None):
    """Madde imli (bullet) paragraf; istege bagli koyu bas ifade."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    return p


def add_figure(doc, filename, caption, width=FIG_W):
    """Ortalanmis bir figur + altina 'Sekil N' aciklamasi ekler."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.name = "Times New Roman"
    cr.font.color.rgb = DARK


def add_table(doc, headers, rows, caption=None, col_widths=None,
              align_right_from=1):
    """
    Basligi golgeli, kenarlikli, satirlari almasik renkli bir tablo ekler.
    headers: sutun basliklari (list[str])
    rows   : satirlar (list[list[str]])
    """
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(2)
        cr = cap.add_run(caption)
        cr.bold = True
        cr.font.size = Pt(10.5)
        cr.font.name = "Times New Roman"
        cr.font.color.rgb = DARK

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    # Baslik satiri
    hdr = table.rows[0].cells
    set_repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_background(hdr[i], HEAD_FILL)
        cp = hdr[i].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.space_before = Pt(2)
        r = cp.add_run(str(h))
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10.5)
        r.font.name = "Times New Roman"

    # Veri satirlari
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                set_cell_background(cells[ci], ALT_FILL)
            cp = cells[ci].paragraphs[0]
            cp.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci < align_right_from
                            else WD_ALIGN_PARAGRAPH.CENTER)
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.line_spacing = 1.0
            r = cp.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = "Times New Roman"

    # Sutun genislikleri
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# --------------------------------------------------------------------------- #
# KAPAK SAYFASI
# --------------------------------------------------------------------------- #
def build_cover(doc):
    para(doc, "ANKARA ÜNİVERSİTESİ", align=WD_ALIGN_PARAGRAPH.CENTER,
         bold=True, size=15, color=BLUE, space_before=18, space_after=2)
    para(doc, "MÜHENDİSLİK FAKÜLTESİ", align=WD_ALIGN_PARAGRAPH.CENTER,
         bold=True, size=13, color=DARK, space_after=2)
    para(doc, "YAZILIM MÜHENDİSLİĞİ BÖLÜMÜ", align=WD_ALIGN_PARAGRAPH.CENTER,
         bold=True, size=13, color=DARK, space_after=18)

    para(doc, "YMH340 — VERİ MADENCİLİĞİ", align=WD_ALIGN_PARAGRAPH.CENTER,
         bold=True, size=13, color=DARK, space_after=2)
    para(doc, "2025–2026 Bahar Dönemi  •  Dönem Projesi Raporu",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=DARK, space_after=40)

    # Cizgi
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), HEAD_FILL)
    pbdr.append(bottom)
    line._p.get_or_add_pPr().append(pbdr)

    para(doc, "Avrupa Futbol Liglerinde Maç Sonuçlarının Makine Öğrenmesi "
              "Algoritmaları ve İddia Oranları Kullanılarak Tahmin Edilmesi",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=17, color=BLUE,
         space_before=30, space_after=8, line_spacing=1.3)

    para(doc, "Çok Sınıflı Sınıflandırma Yaklaşımı ile Spor Analitiği",
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=DARK,
         space_after=44)

    # Ekip uyeleri tablosu
    para(doc, "Ekip Üyeleri", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
         size=12.5, color=DARK, space_after=4)
    members = [
        ["1", "25290085", "Ata Sakık"],
        ["2", "23291269", "Taha Furkan Tosun"],
        ["3", "—", "Yunus Emre Arslan"],
        ["4", "—", "Enes Ayabakan"],
    ]
    t = add_table(doc, ["#", "Öğrenci No", "Adı – Soyadı"], members,
                  col_widths=[Cm(1.2), Cm(4.0), Cm(7.5)], align_right_from=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    para(doc, "Teslim Tarihi: Haziran 2026", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=11, color=DARK, space_before=40)
    para(doc, "Veri madenciliği görevleri Python (pandas, scikit-learn, "
              "matplotlib/seaborn) ortamında gerçekleştirilmiştir.",
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, color=DARK)

    doc.add_page_break()


# --------------------------------------------------------------------------- #
# TABLO VERILERINI OKU
# --------------------------------------------------------------------------- #
def load_tables():
    t = {}
    t["pivot"] = pd.read_csv(TAB / "06_dogruluk_pivot.csv")
    t["clf"] = pd.read_csv(TAB / "07_siniflandirma_raporu.csv")
    t["miss"] = pd.read_csv(TAB / "01_eksik_veri_raporu.csv")
    t["out"] = pd.read_csv(TAB / "02_aykiri_deger_raporu.csv")
    t["imp"] = pd.read_csv(TAB / "04_ozellik_onem_skorlari.csv")
    t["full"] = pd.read_csv(TAB / "05_model_sonuclari_tum.csv")
    return t


# --------------------------------------------------------------------------- #
# BOLUM 1 — GIRIS VE ILGILI CALISMALAR
# --------------------------------------------------------------------------- #
def section1(doc):
    heading(doc, "1. Giriş ve İlgili Çalışmalar")

    para(doc,
        "Futbol, dünya genelinde en çok takip edilen spor dalı olmakla birlikte, "
        "etrafında oluşan dev bahis ekonomisi ve kulüplerin veriye dayalı karar "
        "alma ihtiyacı nedeniyle spor analitiğinin en verimli uygulama "
        "alanlarından biridir. Bu projenin amacı, Avrupa’nın yirmi beş farklı "
        "liginde oynanmış yaklaşık yüz iki bin maça ait geçmiş verileri, bahis "
        "(iddaa) oranlarını ve maç içi takım istatistiklerini (şut, isabetli şut, "
        "korner, faul, kart vb.) kullanarak bir maçın sonucunu — Ev Sahibi "
        "Galibiyeti, Beraberlik veya Deplasman Galibiyeti — tahmin eden bir "
        "sınıflandırma modeli geliştirmektir. Çalışma ayrıca, makine öğrenmesi "
        "yaklaşımlarının bahis piyasasının ima ettiği olasılıklara kıyasla ek bir "
        "öngörü gücü sağlayıp sağlamadığını ve futbol tahmininde bilinen "
        "“beraberlik sınıfının tahmin edilmesindeki zorluk” olgusunu incelemeyi "
        "hedeflemektedir.",
        first_line=0.75)

    para(doc,
        "Problemin önemi iki yönlüdür. Yöntemsel açıdan, üç sınıflı ve dengesiz "
        "dağılıma sahip (ev sahibi galibiyetlerinin baskın olduğu) bir "
        "sınıflandırma problemi olması; veri madenciliğinde eksik veri yönetimi, "
        "özellik mühendisliği, özellik seçimi ve karşılaştırmalı model "
        "değerlendirmesi gibi temel adımların tamamını gerçek ve gürültülü bir "
        "veri seti üzerinde uygulamaya olanak tanır. Uygulama açısından ise "
        "elde edilen bulgular, kulüplerin maç analizinden bahis piyasası "
        "verimliliğinin değerlendirilmesine kadar geniş bir yelpazede karşılık "
        "bulmaktadır.",
        first_line=0.75)

    para(doc,
        "Literatürde bu probleme yönelik çok sayıda çalışma bulunmaktadır. "
        "Öneri raporumuzda da belirtilen iki temel çalışma aşağıda "
        "özetlenmiştir:", first_line=0.75)

    para(doc,
        "[1] numaralı çalışmada, İngiltere Premier "
        "Lig maç sonuçlarını tahmin etmek amacıyla takım istatistikleri ve oyuncu "
        "metriklerinden oluşan bir veri seti kullanılmıştır. Veri ön işleme "
        "aşamasında eksik veriler temizlenmiş ve hedef değişkeni en çok etkileyen "
        "öznitelikler algoritmik olarak (özellik seçimi ile) belirlenmiştir. "
        "Sınıflandırma aşamasında Random Forest, XGBoost ve LightGBM "
        "algoritmaları karşılaştırılmış; testler sonucunda Random Forest modeli "
        "%87,14 doğruluk oranıyla en başarılı yöntem olarak öne çıkmıştır. Bu "
        "yüksek doğruluk, tek bir lige odaklanılması ve zengin oyuncu düzeyi "
        "metriklerinin kullanılmasıyla ilişkilendirilebilir.",
        first_line=0.75)

    para(doc,
        "İkinci çalışmada [2] ise maç sonuçlarını tahmin etmek için geçmiş maç "
        "istatistikleri, FIFA takım derecelendirmeleri ve iddia oranları "
        "birleştirilerek, çift değişkenli Poisson dağılımını da içeren hibrit bir "
        "makine öğrenmesi çerçevesi önerilmiştir. Çalışmada beraberlik (draw) "
        "sınıfını tahmin etmenin diğer sınıflara kıyasla belirgin biçimde zor "
        "olduğu vurgulanmıştır. Buna rağmen uygulanan Random Forest modeli %56,25 "
        "genel doğruluk oranı sağlayarak, bahis oranlarının ve geçmiş verilerin "
        "tahmin modellerine başarıyla entegre edilebileceğini göstermiştir. Bu "
        "çalışmanın kurgusu (oranlar + istatistikler, üç sınıflı sonuç) projemizle "
        "büyük ölçüde örtüştüğü için sonuçlarımız en çok bu çalışma ile "
        "karşılaştırılmıştır.",
        first_line=0.75)

    para(doc,
        "Bu iki çalışma, hem özellik seçiminin hem de bahis oranları ile maç "
        "istatistiklerinin birlikte kullanılmasının değerini ortaya koymaktadır. "
        "Projemiz, tek bir lig yerine yirmi beş ligi kapsayan çok daha geniş ve "
        "heterojen bir veri seti üzerinde, altı farklı algoritmayı tüm özellikler "
        "ve seçilmiş özelliklerle karşılaştırarak bu literatüre katkı sağlamayı "
        "amaçlamaktadır.",
        first_line=0.75)


# --------------------------------------------------------------------------- #
# BOLUM 2 — PROBLEM TANIMI VE ALGORITMALAR
# --------------------------------------------------------------------------- #
def section2(doc):
    heading(doc, "2. Problem Tanımı ve Veri Madenciliği Algoritmaları")

    heading(doc, "2.1. Görev Tanımı", level=2)
    para(doc,
        "Bu proje, denetimli öğrenme (supervised learning) kapsamında yer alan "
        "çok sınıflı bir sınıflandırma (multi-class classification) problemidir. "
        "Her gözlem (maç) için hedef değişken, ev sahibi ve deplasman takımlarının "
        "attığı gol sayıları karşılaştırılarak türetilen üç ayrık sınıftan birini "
        "alır:", first_line=0.75)
    bullet(doc, "ev sahibi attığı golün fazla olduğu maçlar.", "Sınıf 1 — Ev Sahibi Galibiyeti: ")
    bullet(doc, "iki takımın gol sayısının eşit olduğu maçlar.", "Sınıf 0 — Beraberlik: ")
    bullet(doc, "deplasman takımının attığı golün fazla olduğu maçlar.", "Sınıf 2 — Deplasman Galibiyeti: ")
    para(doc,
        "Hedef değişken sürekli değil kategorik (kesikli) olduğundan problem bir "
        "regresyon değil sınıflandırma problemidir; etiketli (sonucu bilinen) "
        "geçmiş maçlarla öğrenildiği için de denetimsiz (kümeleme) değil "
        "denetimli öğrenmedir. Modelin başarısı, hiçbir bilgi kullanmadan her zaman "
        "en sık görülen sınıfı (Ev Sahibi Galibiyeti) tahmin eden “çoğunluk-sınıf "
        "taban çizgisi” (baseline) ile karşılaştırılarak değerlendirilmiştir.",
        first_line=0.75)

    heading(doc, "2.2. Seçilen Veri Madenciliği Algoritmaları", level=2)
    para(doc,
        "Karşılaştırma için, farklı öğrenme paradigmalarını temsil eden altı "
        "sınıflandırma algoritması seçilmiştir (proje önerisinde belirtilen Karar "
        "Ağacı, Random Forest ve Naive Bayes zorunlu olarak dâhil edilmiş, "
        "kıyaslamayı zenginleştirmek için üç algoritma daha eklenmiştir):",
        first_line=0.75)
    bullet(doc, "Özellik uzayını eksen-hizalı bölmelerle ardışık olarak ayıran, "
                "yorumlanabilirliği yüksek bir ağaç modelidir. Tek başına aşırı "
                "öğrenmeye (overfitting) eğilimlidir.", "Karar Ağacı (Decision Tree): ")
    bullet(doc, "Çok sayıda karar ağacını rastgelelik (bagging + öznitelik "
                "örnekleme) ile birleştiren bir topluluk (ensemble) yöntemidir; "
                "tek ağaca göre daha kararlı ve doğru sonuç verir.",
                "Rastgele Orman (Random Forest): ")
    bullet(doc, "Bayes teoremine dayanan, özelliklerin sınıf koşulunda bağımsız "
                "olduğunu varsayan olasılıksal bir modeldir; çok hızlıdır ve "
                "temel bir karşılaştırma ölçütü sunar.", "Naive Bayes (Gaussian): ")
    bullet(doc, "Sınıf olasılıklarını doğrusal bir sınırla modelleyen, "
                "katsayıları yorumlanabilir istatistiksel bir yöntemdir.",
                "Lojistik Regresyon (Logistic Regression): ")
    bullet(doc, "Bir örneği, özellik uzayında kendisine en yakın K komşusunun "
                "çoğunluk sınıfına atayan örnek-tabanlı (lazy) bir yöntemdir; "
                "ölçeklemeye duyarlıdır.", "K-En Yakın Komşu (KNN): ")
    bullet(doc, "Ağaçları, önceki ağaçların hatalarını düzeltecek şekilde "
                "ardışık (boosting) ekleyen güçlü bir topluluk yöntemidir; "
                "[1] ve [2]’deki XGBoost/LightGBM ailesinin scikit-learn "
                "karşılığıdır.", "Gradient Boosting (HistGradientBoosting): ")
    para(doc,
        "Ölçeklemeye duyarlı modeller (Lojistik Regresyon, KNN, Naive Bayes) için "
        "öznitelikler bir işlem hattı (Pipeline) içinde StandardScaler ile "
        "standartlaştırılmış; ağaç tabanlı modeller (Karar Ağacı, Random Forest, "
        "Gradient Boosting) ölçekleme gerektirmediğinden doğrudan eğitilmiştir.",
        first_line=0.75)


# --------------------------------------------------------------------------- #
# BOLUM 3 — VERI OZETI
# --------------------------------------------------------------------------- #
def section3(doc, t):
    heading(doc, "3. Veri Özeti")
    para(doc,
        "Veri Kaynağı. Çalışmada kullanılan veri seti, ekibimiz tarafından farklı "
        "açık kaynaklardan (yerel lig arşivleri ve Wikipedia maç kayıtları) "
        "derlenmiş olan “Avrupa Futbol Veritabanı” "
        "(EUROPEAN_FOOTBALL_DATABASE_FULL.csv) dosyasıdır. Veri seti 102.815 maç "
        "(satır) ve 28 değişkenden (sütun) oluşmakta; 2000–2001’den 2025–2026’ya "
        "kadar 26 sezonu ve Premier League, La Liga, Serie A, Bundesliga, Süper "
        "Lig dâhil 25 farklı lig/turnuvayı kapsamaktadır. Dosya boyutunun "
        "büyüklüğü (≈14 MB, bellekte ≈58 MB) ve değişken sayısı nedeniyle WEKA ve "
        "Orange gibi masaüstü araçlar performans açısından zorlandığından, tüm "
        "analiz Python ortamında gerçekleştirilmiştir.",
        first_line=0.75)

    para(doc,
        "Özellikler (Features). Değişkenler işlevlerine göre dört gruba "
        "ayrılabilir. Aşağıdaki tabloda gruplar ve örnek sütunlar özetlenmiştir.",
        first_line=0.75)

    add_table(doc,
        ["Değişken Grubu", "Örnek Sütunlar", "Açıklama"],
        [
            ["Kimlik / Meta veri", "date, league, season, source, referee",
             "Maçın tarihi, ligi, sezonu ve kaynağı"],
            ["Gol bilgisi", "home_goals, away_goals, ht_home_goals, ht_away_goals",
             "Tam ve ilk yarı gol sayıları (hedef bunlardan türetilir)"],
            ["Bahis oranları", "odds_home, odds_draw, odds_away, odds_over/under_25",
             "Maç öncesi piyasa beklentisi"],
            ["Maç istatistikleri", "home/away_shots, shots_target, corners, fouls, yellow, red",
             "Sahadaki performans göstergeleri"],
        ],
        caption="Tablo 1. Veri setindeki değişken gruplarının özeti.",
        col_widths=[Cm(3.5), Cm(6.0), Cm(6.0)])

    para(doc,
        "Sınıf Niteliği (Hedef Değişken). Veri setinde hazır bir “sonuç” sütunu "
        "bulunmadığından hedef değişken, home_goals ve away_goals "
        "karşılaştırılarak türetilmiştir (1: Ev, 0: Beraberlik, 2: Deplasman). "
        "Sınıf dağılımı dengesizdir: ev sahibi galibiyetleri %45,8 ile baskındır; "
        "bu durum futbolda iyi bilinen “ev sahibi avantajı” olgusunu doğrular ve "
        "modellerin aşması gereken çoğunluk-sınıf taban doğruluğunu (≈%45,8) "
        "belirler. Şekil 1 hedef dağılımını, Şekil 2 ise maç sayısına göre en "
        "büyük ligleri göstermektedir.",
        first_line=0.75)

    add_figure(doc, "01_sinif_dagilimi.png",
               "Şekil 1. Hedef değişkenin (maç sonucu) sınıf dağılımı.")
    add_figure(doc, "02_lig_dagilimi.png",
               "Şekil 2. En çok maça sahip 15 lig.", width=Cm(13.5))


# --------------------------------------------------------------------------- #
# BOLUM 4 — VERI HAZIRLAMA
# --------------------------------------------------------------------------- #
def section4(doc, t):
    heading(doc, "4. Veri Hazırlama")

    heading(doc, "4.1. Eksik Veri Analizi ve Temizleme", level=2)
    para(doc,
        "Eksik değerler, pandas kütüphanesinin isna() fonksiyonuyla sütun bazında "
        "tespit edilmiştir. Veri seti farklı kaynakların birleşiminden oluştuğu "
        "için eksiklik oranları sütunlar arasında büyük farklılık göstermektedir "
        "(Tablo 2, Şekil 3). Hakem (referee) bilgisi maçların yaklaşık %84’ünde, "
        "2,5 alt/üst gol oranları ise %73’ünde eksiktir; maç istatistikleri "
        "(şut, korner, faul, kart) yaklaşık %42 oranında, bahis oranları ise "
        "%22 oranında eksiktir. Buna karşın hedefi belirleyen gol sütunlarında "
        "hiç eksik değer yoktur.",
        first_line=0.75)

    miss = t["miss"]
    miss = miss[miss["eksik_yuzde"] > 0].head(8)
    rows = [[r["sutun"], f"{int(r['eksik_sayi']):,}", f"%{r['eksik_yuzde']:.2f}"]
            for _, r in miss.iterrows()]
    add_table(doc, ["Sütun", "Eksik Kayıt", "Eksik Oran"], rows,
              caption="Tablo 2. En yüksek eksiklik oranına sahip sütunlar.",
              col_widths=[Cm(6.0), Cm(4.5), Cm(4.5)])

    add_figure(doc, "03_eksik_veri.png",
               "Şekil 3. Sütun bazlı eksik veri oranları.", width=Cm(14))

    para(doc,
        "Uygulanan temizleme stratejisi şu şekildedir:", first_line=0.75)
    bullet(doc, "Aşırı eksik ve/veya kimlik niteliğindeki sütunlar (referee, "
                "odds_over_25, odds_under_25, date, home_team, away_team, source) "
                "modellemeden çıkarılmıştır.", "Sütun eleme: ")
    bullet(doc, "İlk yarı gol sütunları (ht_home_goals, ht_away_goals) nihai "
                "skorun bir parçası olduğundan kısmi bir veri sızıntısı (leakage) "
                "riski taşır; bu nedenle özellik kümesinden çıkarılmıştır.",
                "Sızıntı önleme: ")
    bullet(doc, "Modelleme için gerekli olan oran ve maç istatistiği sütunlarında "
                "eksik değeri bulunan satırlar elenmiş (tam-kayıt/complete-case "
                "yaklaşımı), geriye 55.362 eksiksiz maç kalmıştır. Bu kayıtlar "
                "için %42 gibi yüksek oranlı bir atama (imputation) ham sinyali "
                "bozacağından; medyan ile atama, daha küçük ama daha güvenilir bir "
                "tam-kayıt kümesine tercih edilmemiştir.", "Satır eleme: ")
    bullet(doc, "Hedefi doğrudan belirleyen home_goals ve away_goals sütunları, "
                "klasik bir veri sızıntısı kaynağı olduğundan özellik olarak "
                "kesinlikle kullanılmamıştır.", "Hedef sızıntısının engellenmesi: ")

    heading(doc, "4.2. Aykırı Değer (Outlier) Analizi", level=2)
    para(doc,
        "Aykırı değerler, IQR (Çeyrekler Açıklığı) yöntemiyle tespit edilmiştir: "
        "bir değer [Q1 − 1,5×IQR , Q3 + 1,5×IQR] aralığının dışındaysa aykırı "
        "kabul edilir. Şut, korner ve faul gibi sayısal istatistiklerde aykırı "
        "oranı genellikle %3’ün altındadır (Şekil 4, Tablo 3). Kırmızı kart "
        "sütunlarındaki yüksek oran ise, değerlerin medyanının sıfır olması ve "
        "her kırmızı kartın istatistiksel olarak “aykırı” sayılması nedeniyle "
        "yapaydır. Bu aykırı değerler hatalı kayıt değil, futbolun doğal "
        "uç durumları (örneğin 40+ şut atılan bir maç) olduğundan silinmemiş, "
        "yalnızca raporlanmıştır; ağaç tabanlı modeller bu tür değerlere zaten "
        "dayanıklıdır.",
        first_line=0.75)

    out = t["out"]
    sel = out[out["sutun"].isin(
        ["home_shots", "away_shots", "home_corners", "away_fouls",
         "away_red", "odds_home"])]
    rows = [[r["sutun"], f"{r['alt_sinir']:.1f}", f"{r['ust_sinir']:.1f}",
             f"{int(r['aykiri_sayi']):,}", f"%{r['aykiri_yuzde']:.2f}"]
            for _, r in sel.iterrows()]
    add_table(doc, ["Sütun", "Alt Sınır", "Üst Sınır", "Aykırı Sayı", "Aykırı Oran"],
              rows, caption="Tablo 3. Seçili sütunlar için IQR aykırı değer özeti.",
              col_widths=[Cm(4.2), Cm(2.6), Cm(2.6), Cm(3.0), Cm(2.6)])

    add_figure(doc, "04_aykiri_deger_boxplot.png",
               "Şekil 4. Maç istatistikleri için aykırı değer kutu grafikleri.")

    heading(doc, "4.3. Özellik Mühendisliği ve Korelasyon", level=2)
    para(doc,
        "Ham değişkenlere ek olarak, maç sonucunu daha güçlü açıklayan altı yeni "
        "özellik türetilmiştir. (i) İma edilen olasılıklar: her bahis oranının "
        "tersi (1/oran) alınıp üç sonucun toplamı 1 olacak şekilde normalize "
        "edilerek bahis ofisi kâr marjı (overround) giderilmiş ve piyasanın maça "
        "dair beklentisi olasılık biçiminde ifade edilmiştir. (ii) Fark "
        "özellikleri: ev ve deplasman istatistiklerinin farkı (şut farkı, "
        "isabetli şut farkı, korner farkı) hesaplanarak sahadaki göreli üstünlük "
        "tek bir değerde özetlenmiştir. Böylece özellik sayısı 15 ham + 6 "
        "türetilmiş = 21’e ulaşmıştır. Şekil 5’teki korelasyon ısı haritası, "
        "isabetli şut farkı ve ima edilen olasılıkların hedefle en güçlü ilişkiye "
        "sahip değişkenler olduğunu göstermektedir.",
        first_line=0.75)

    add_figure(doc, "05_korelasyon_haritasi.png",
               "Şekil 5. Özellikler ve hedef arasındaki Pearson korelasyon "
               "ısı haritası.", width=Cm(15))

    heading(doc, "4.4. Özellik Seçimi (Feature Selection)", level=2)
    para(doc,
        "Özellik seçimi için birbirini tamamlayan iki yöntem kullanılmıştır: "
        "(1) ANOVA F-testi (SelectKBest, f_classif) ile her özelliğin sınıflar "
        "arasındaki ayırt ediciliği tek değişkenli olarak ölçülmüş; (2) Random "
        "Forest öznitelik önem skorları ile değişkenler arası etkileşimleri de "
        "yakalayan model-tabanlı bir sıralama elde edilmiştir. İki skor [0–1] "
        "aralığına normalize edilip ortalaması alınarak birleşik bir önem skoru "
        "oluşturulmuş ve en yüksek skorlu ilk 10 özellik seçilmiştir (Şekil 6). "
        "Seçilen özellikler şunlardır: isabetli şut farkı, ima edilen ev/deplasman "
        "galibiyet olasılıkları, ev sahibi/deplasman oranları, ev ve deplasman "
        "isabetli şut sayıları, şut farkı, ima edilen beraberlik olasılığı ve "
        "deplasman şut sayısı. Kart ve faul gibi değişkenlerin önem skorlarının "
        "düşük olması, bunların maç sonucu üzerinde sınırlı doğrudan etkiye sahip "
        "olduğunu göstermektedir.",
        first_line=0.75)

    add_figure(doc, "07_ozellik_onem.png",
               "Şekil 6. ANOVA F ve Random Forest skorlarının birleşimiyle "
               "elde edilen özellik önem sıralaması (ilk 10 seçilmiştir).",
               width=Cm(14.5))


# --------------------------------------------------------------------------- #
# BOLUM 5 — SONUCLAR
# --------------------------------------------------------------------------- #
def section5(doc, t):
    heading(doc, "5. Sonuçlar")

    heading(doc, "5.1. Model Seçimi ve Eğitimi", level=2)
    para(doc,
        "55.362 maçlık temizlenmiş veri seti, sınıf oranları korunarak (stratified) "
        "%80 eğitim (44.289 maç) ve %20 test (11.073 maç) olarak ikiye "
        "ayrılmıştır. Altı algoritmanın her biri iki ayrı özellik kümesiyle "
        "eğitilmiştir: (a) tüm 21 özellik ve (b) özellik seçimiyle belirlenen 10 "
        "özellik. Modellerin kararlılığı 5 katlı çapraz doğrulama (5-fold "
        "cross-validation) ile de teyit edilmiş; tekrarlanabilirlik için tüm "
        "rastgelelik tohumu sabitlenmiştir (random_state = 42). Başarı ölçütü "
        "olarak doğruluk (accuracy) yanında, dengesiz sınıf yapısı nedeniyle "
        "makro-ortalama F1 skoru da raporlanmıştır.",
        first_line=0.75)

    heading(doc, "5.2. Karşılaştırmalı Doğruluk Sonuçları", level=2)
    para(doc,
        "Tablo 4, altı algoritmanın tüm özellikler ve seçilmiş özelliklerle elde "
        "ettiği test doğruluklarını ve aradaki farkı; Şekil 7 ise aynı sonuçların "
        "görsel karşılaştırmasını sunmaktadır. Tüm modeller, %44,78’lik "
        "çoğunluk-sınıf taban doğruluğunu belirgin biçimde aşmıştır.",
        first_line=0.75)

    piv = t["pivot"].sort_values("Tüm Özellikler", ascending=False)
    rows = []
    for _, r in piv.iterrows():
        rows.append([
            r["Algoritma"],
            f"%{r['Tüm Özellikler']*100:.2f}",
            f"%{r['Seçilmiş Özellikler']*100:.2f}",
            f"{r['Fark']*100:+.2f} puan",
        ])
    add_table(doc,
        ["Algoritma", "Tüm Özellikler (21)", "Seçilmiş Özellikler (10)", "Fark"],
        rows,
        caption="Tablo 4. Algoritmaların doğruluk sonuçları: tüm özellikler ile "
                "seçilmiş özelliklerin karşılaştırması.",
        col_widths=[Cm(4.5), Cm(4.0), Cm(4.3), Cm(2.7)])

    add_figure(doc, "08_dogruluk_karsilastirma.png",
               "Şekil 7. Algoritmaların test doğruluğu: tüm vs seçilmiş özellikler.")

    para(doc,
        "En yüksek doğruluğu %62,11 ile Gradient Boosting elde etmiş; onu %61,41 "
        "ile Lojistik Regresyon ve %60,96 ile Random Forest izlemiştir. Naive "
        "Bayes (%52,86) ve tek Karar Ağacı (%56,94) en düşük başarıyı göstermiştir. "
        "Dikkat çekici bir bulgu, özellik sayısının 21’den 10’a düşürülmesine "
        "rağmen doğruluğun çok az değişmesidir (en güçlü modellerde 1–4 puan): "
        "bu, seçilen 10 özelliğin bilginin büyük kısmını taşıdığını ve özellik "
        "seçiminin model karmaşıklığını/eğitim süresini azaltırken yorumlanabilirliği "
        "artırdığını göstermektedir. Lojistik Regresyon ve Karar Ağacı gibi bazı "
        "modellerde seçilmiş özelliklerle neredeyse aynı veya daha iyi sonuç "
        "alınması, gürültülü değişkenlerin elenmesinin olumlu etkisine işaret eder.",
        first_line=0.75)

    heading(doc, "5.3. En İyi Modelin Ayrıntılı Değerlendirmesi", level=2)
    para(doc,
        "En başarılı model olan Gradient Boosting için sınıf bazında "
        "kesinlik (precision), duyarlılık (recall) ve F1 skorları Tablo 5’te, "
        "karmaşıklık matrisi ise Şekil 8’de verilmiştir.",
        first_line=0.75)

    clf = t["clf"]
    label_map = {"Ev Sahibi Galibiyeti": "Ev Sahibi Galibiyeti (1)",
                 "Beraberlik": "Beraberlik (0)",
                 "Deplasman Galibiyeti": "Deplasman Galibiyeti (2)"}
    rows = []
    for _, r in clf.iterrows():
        name = str(r["sinif"])
        if name in label_map:
            rows.append([label_map[name], f"{r['precision']:.3f}",
                         f"{r['recall']:.3f}", f"{r['f1-score']:.3f}",
                         f"{int(r['support']):,}"])
    add_table(doc,
        ["Sınıf", "Kesinlik", "Duyarlılık", "F1 Skoru", "Örnek Sayısı"], rows,
        caption="Tablo 5. Gradient Boosting modelinin sınıf bazlı başarı metrikleri "
                "(test kümesi).",
        col_widths=[Cm(5.5), Cm(2.7), Cm(2.7), Cm(2.5), Cm(2.6)])

    add_figure(doc, "09_karmasiklik_matrisi.png",
               "Şekil 8. Gradient Boosting modelinin karmaşıklık matrisi "
               "(solda sayılar, sağda satıra göre normalize oranlar).")

    para(doc,
        "Sonuçlar, sınıflar arasında belirgin bir başarı farkı olduğunu ortaya "
        "koymaktadır. Model, ev sahibi galibiyetlerini %81,8 ve deplasman "
        "galibiyetlerini %68,2 duyarlılıkla doğru tahmin ederken; beraberlikleri "
        "yalnızca %19,9 duyarlılıkla bilebilmektedir. Karmaşıklık matrisi, "
        "gerçekte beraberlikle biten maçların büyük bölümünün model tarafından ev "
        "sahibi veya deplasman galibiyeti olarak sınıflandırıldığını "
        "göstermektedir. Bu durum, beraberlik sınıfının hem azınlıkta olması hem "
        "de istatistiksel olarak diğer iki sınıfın “arasında” kalması nedeniyle "
        "literatürde [2] de vurgulanan iyi bilinen bir zorluktur.",
        first_line=0.75)

    add_figure(doc, "06_oran_olasilik_dagilimi.png",
               "Şekil 9. İma edilen ev sahibi galibiyet olasılığının gerçek maç "
               "sonucuna göre dağılımı.", width=Cm(14))
    para(doc,
        "Şekil 9, neden modellerin işe yaradığını sezgisel olarak açıklar: bahis "
        "piyasasının ev sahibine biçtiği olasılık yükseldikçe gerçek sonuç ev "
        "sahibi galibiyetine, düştükçe deplasman galibiyetine kaymakta; "
        "beraberlikler ise orta bölgede yoğunlaşmaktadır. Bu örtüşme bölgesi, "
        "beraberlik tahminindeki güçlüğün de görsel kanıtıdır.",
        first_line=0.75)


# --------------------------------------------------------------------------- #
# BOLUM 6 — TARTISMA VE SONUC
# --------------------------------------------------------------------------- #
def section6(doc):
    heading(doc, "6. Tartışma ve Sonuç")

    heading(doc, "6.1. Çalışmanın Önemi ve Katkıları", level=2)
    para(doc,
        "Bu çalışma, 25 ligi ve 26 sezonu kapsayan geniş ve gürültülü gerçek bir "
        "veri seti üzerinde, eksik veri yönetiminden özellik seçimine ve altı "
        "algoritmanın karşılaştırmalı değerlendirmesine uzanan eksiksiz bir veri "
        "madenciliği hattı sunmaktadır. Projeye en uygun algoritmanın, ardışık "
        "topluluk yöntemi olan Gradient Boosting (%62,11) olduğu; doğrusal bir "
        "model olan Lojistik Regresyon’un (%61,41) ona çok yakın bir başarıyla "
        "güçlü bir alternatif oluşturduğu görülmüştür. Tek Karar Ağacı ve Naive "
        "Bayes ise görece zayıf kalmıştır. Topluluk ve doğrusal modellerin öne "
        "çıkması, problemde hem doğrusal (oranların ima ettiği olasılıklar) hem de "
        "doğrusal olmayan (istatistikler arası etkileşimler) sinyallerin "
        "bulunduğuna işaret eder.",
        first_line=0.75)
    para(doc,
        "Bulgularımız önceki çalışmalarla tutarlıdır. Kurgusu projemizle örtüşen "
        "[2] numaralı çalışmada Random Forest %56,25 doğruluk elde etmişti; "
        "bizim Random Forest sonucumuz (%60,96) ve en iyi modelimiz (%62,11) bu "
        "değerin bir miktar üzerindedir; bunun, maç istatistiklerinin de özellik "
        "olarak kullanılması ve çok daha büyük bir eğitim kümesiyle "
        "açıklanabileceğini değerlendiriyoruz. [1] numaralı çalışmadaki %87,14’lük "
        "yüksek doğruluğa ulaşılamaması ise beklenen bir sonuçtur: o çalışma tek "
        "bir lige odaklanmakta ve zengin oyuncu düzeyi metrikleri kullanmaktadır; "
        "bizim veri setimiz ise çok daha heterojendir. Ayrıca her iki çalışmada da "
        "vurgulanan beraberlik sınıfının tahmin güçlüğü, bizim sonuçlarımızda da "
        "(beraberlik duyarlılığı %19,9) net biçimde gözlenmiştir.",
        first_line=0.75)
    para(doc,
        "Özellik seçimi açısından önemli bir katkı, özellik sayısının yarıdan "
        "fazlasının (21→10) elenmesine rağmen doğruluğun korunmasıdır; bu, "
        "modelin daha sade, hızlı ve yorumlanabilir hâle getirilebileceğini "
        "gösterir. İsabetli şut farkı ve bahis oranlarından türetilen olasılıkların "
        "en belirleyici değişkenler olması da hem futbol sezgisiyle hem de "
        "literatürle uyumludur.",
        first_line=0.75)

    heading(doc, "6.2. Kısıtlamalar ve Gelecek Çalışmalar", level=2)
    para(doc, "Çalışma sırasında karşılaşılan başlıca kısıtlamalar ve gelecekteki "
              "geliştirme önerileri şunlardır:", first_line=0.75)
    bullet(doc, "Maç istatistikleri (şut, korner vb.) ancak maç oynandıktan sonra "
                "bilindiğinden, bu çalışma saf bir “maç öncesi tahmin” değil, maç "
                "verilerine dayalı bir sınıflandırmadır. Gelecekte yalnızca maç "
                "öncesi bilgilerle (oranlar, takım form/Elo puanları) çalışan bir "
                "model kurularak gerçek öngörü gücü ölçülebilir.",
                "Bilgi türü: ")
    bullet(doc, "Eksik veriler nedeniyle kayıtların yaklaşık %46’sı modelleme dışı "
                "kalmıştır. İleri atama (imputation) teknikleri veya istatistiklerin "
                "bulunmadığı maçlar için oran-tabanlı ayrı bir model ile bu kayıplar "
                "azaltılabilir.", "Veri kaybı: ")
    bullet(doc, "Beraberlik sınıfının düşük duyarlılığı en temel kısıttır. SMOTE "
                "gibi yeniden örnekleme, maliyet-duyarlı öğrenme veya sınıf "
                "ağırlıklandırma yöntemleriyle azınlık sınıfının başarısı "
                "artırılabilir.", "Sınıf dengesizliği: ")
    bullet(doc, "Zaman serisi niteliği göz ardı edilip rastgele bölme yapılmıştır; "
                "sezonlara göre kronolojik (temporal) bölme daha gerçekçi bir "
                "değerlendirme sağlar.", "Doğrulama kurgusu: ")
    bullet(doc, "Hiperparametreler büyük ölçüde varsayılan değerlerde "
                "bırakılmıştır; ızgara/rastgele arama (GridSearch) ve XGBoost / "
                "LightGBM gibi modellerle ek kazanım elde edilebilir.",
                "Model ayarı: ")
    para(doc,
        "Sonuç olarak, makine öğrenmesi yöntemleri Avrupa futbol liglerinde maç "
        "sonucunu çoğunluk-sınıf taban çizgisinin belirgin biçimde üzerinde "
        "(≈%62) tahmin edebilmekte; bahis oranlarından türetilen olasılıklar ile "
        "isabetli şut gibi temel maç istatistikleri bu başarının çekirdeğini "
        "oluşturmaktadır. Beraberlik sınıfının tahmini ise hem bu çalışmanın hem "
        "de alanyazının açık problemi olmayı sürdürmektedir.",
        first_line=0.75)


# --------------------------------------------------------------------------- #
# REFERANSLAR (IEEE)
# --------------------------------------------------------------------------- #
def references(doc):
    heading(doc, "Referanslar")
    refs = [
        "[1]  “A Machine Learning Model for Analysis and Prediction of Football "
        "Match Outcomes in the English Premier League,” International Journal of "
        "Research and Innovation in Applied Science (IJRIAS).",
        "[2]  “A Hybrid Machine Learning Framework for Soccer Match Outcome "
        "Prediction: Incorporating Bivariate Poisson Distribution,” ITM Web of "
        "Conferences.",
        "[3]  F. Pedregosa et al., “Scikit-learn: Machine Learning in Python,” "
        "Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.",
        "[4]  W. McKinney, “Data Structures for Statistical Computing in Python,” "
        "in Proc. 9th Python in Science Conference (SciPy), 2010, pp. 56–61.",
        "[5]  L. Breiman, “Random Forests,” Machine Learning, vol. 45, no. 1, "
        "pp. 5–32, 2001.",
        "[6]  J. H. Friedman, “Greedy Function Approximation: A Gradient Boosting "
        "Machine,” Annals of Statistics, vol. 29, no. 5, pp. 1189–1232, 2001.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        run = p.add_run(r)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"


# --------------------------------------------------------------------------- #
# ANA URETIM
# --------------------------------------------------------------------------- #
def main():
    doc = Document()
    style_normal(doc)

    # Sayfa boyutu A4 + 2.5 cm kenar bosluklari
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    build_cover(doc)
    add_page_numbers(doc.sections[0])

    t = load_tables()
    section1(doc)
    section2(doc)
    section3(doc, t)
    section4(doc, t)
    section5(doc, t)
    section6(doc)
    references(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Rapor olusturuldu -> {OUT}")
    print(f"Toplam paragraf sayisi: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
